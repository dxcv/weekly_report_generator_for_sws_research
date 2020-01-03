import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
# from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import xlwings as xw
from openpyxl import load_workbook
import tkinter as tk
from tkinter import filedialog
from copy import deepcopy
import tushare as ts
import pandas as pd
import math
import re
from datetime import datetime, timedelta
from WindPy import *
w.start()


def generate_sections():
    root = tk.Tk()
    root.withdraw()
    print("运行本程序前建议将相关数据备份")
    print("------------------------- START -------------------------")
    date = input("请输入本周最后一个交易日日期（YYYYMMDD，例：20190816）：")
    print("选择需要修改的【周报】文件")
    file_path_2 = filedialog.askopenfilename()
    doc_2 = docx.Document(file_path_2)

    def cycle_letter(arr, level):
        tempArr = []
        letterArr = [chr(i) for i in range(65, 91)]
        arrNum = len(arr)
        if level == 0 or arrNum == 0:
            return letterArr
        for index in range(arrNum):
            for letter in letterArr:
                tempArr.append(arr[index] + letter)
        return tempArr

    def generate_excel_col_name(num):
        tempVal = 1
        level = 1
        while tempVal:
            tempVal = num / (math.pow(26, level))
            if tempVal > 1:
                level += 1
            else:
                break

        excelArr = []
        tempArr = []
        for index in range(level):
            tempArr = cycle_letter(tempArr, index)
            for numIndex in range(len(tempArr)):
                if (len(excelArr) < num):
                    excelArr.append(tempArr[numIndex])
                else:
                    return excelArr

    characters = generate_excel_col_name(300)

    # Section 1: 生成一周板块回顾
    print("-----------------Step 1: 生成一周板块回顾-----------------")
    print("导入【本周投资提示】表格，导入前请确认各sheet中【数据提取日期】单元格已更改为本周最后一个交易日且数据已更新")
    file_path_3 = filedialog.askopenfilename()
    print("本周投资提示表格已导入")

    def generate_section_1(file_path, doc_2):

        def check_data(date, sheet_title, sheet, start_no, end_no):
            """
            检查本周投资提示表格内数据是否更新
            """
            data_range = sheet["B" + str(start_no) + ":B" + str(end_no)].value
            data_range_str = ""
            for i in range(len(data_range)):
                if i != len(data_range):
                    data_range_str += str(data_range[i]) + ","
                else:
                    data_range_str += str(data_range[i])

            print("正在检查" + sheet_title + "数据是否已更新......")
            data = w.wss(data_range_str, "pct_chg", "tradeDate=" + date + ";cycle=W")
            df = pd.DataFrame(data.Data).T
            df.index = data.Codes
            df.columns = data.Fields

            check_df = pd.DataFrame(sheet_1["D" + str(start_no) + ":D" + str(end_no)].value,
                                    index=sheet_1["B" + str(start_no) + ":B" + str(end_no)].value, columns=data.Fields)

            check_list = []
            for i in range(len(df)):
                for j in range(len(check_df)):
                    if df.index[i] == check_df.index[j]:
                        if df.iloc[i, 0] == check_df.iloc[j, 0]:
                            check_list.append(1)
                        else:
                            check_list.append(0)
            if all(check_list):
                print(sheet_title + "数据已更新")
            else:
                for i in range(len(check_list)):
                    if check_list[i] == 0:
                        print(str(df.index[i]) + "数据不匹配，请更新后重试！")
                        exit()

        workbook = xw.Book(file_path)
        sheet_1 = workbook.sheets("大盘情况和子行业排名")
        sheet_2 = workbook.sheets("白酒等指数")
        sheet_3 = workbook.sheets("个股排名")
        sheet_4 = workbook.sheets["一周板块回顾"]

        print("正在检查数据提取日期是否已更新......")
        date_datetime = datetime.strptime(date, "%Y%m%d")
        if sheet_4["B1"].value == date_datetime:
            pass
        else:
            print("【一周板块回顾】sheet日期未更新，请更新后重试！")
            exit()
        if sheet_1["B1"].value == date_datetime:
            pass
        else:
            print("【大盘情况和子行业排名】sheet日期未更新，请更新后重试！")
            exit()
        if sheet_2["B1"].value == date_datetime:
            pass
        else:
            print("【白酒等指数】sheet日期未更新，请更新后重试！")
            exit()
        if sheet_3["B1"].value == date_datetime:
            pass
        else:
            print("【个股排名】sheet日期未更新，请更新后重试！")
            exit()

        # date = input("请输入本周最后一个交易日日期（YYYYMMDD，例：20190816）：")
        check_data(date, "大盘情况和子行业排名", sheet_1, 7, 36)
        check_data(date, "白酒等指数", sheet_2, 4, 12)
        check_data(date, "个股排名", sheet_3, 4, 97)

        print("正在执行修改......")

        def generate_part_1():

            spyl = sheet_1["D27"].value
            if spyl >= 0:
                str_1 = "上涨"
            else:
                str_1 = "下跌"
            spyl_str = "上周食品饮料板块" + str_1 + format(abs(spyl), ".2f") + "%，"

            bj = sheet_2["D11"].value
            if bj >= 0:
                str_2 = "上涨"
            else:
                str_2 = "下跌"
            bj_str = "其中白酒" + str_2 + format(abs(bj), ".2f") + "%，"

            szzz = sheet_1["D19"].value
            if szzz >= 0:
                str_3 = "上涨"
            else:
                str_3 = "下跌"
            szzz_str = "上证综指" + str_3 + format(abs(szzz), ".2f") + "%，"

            active_return = round(spyl, 2) - round(szzz, 2)
            if active_return >= 0:
                str_4 = "跑赢"
            else:
                str_4 = "跑输"
            active_return_str = str_4 + "大盘" + format(abs(active_return), ".2f") + "pct，"

            index = sheet_1["A7:A36"].value
            pctchange = sheet_1["D7:D36"].value
            pctchange_df = pd.DataFrame(pctchange, index=index)
            pctchange_df = pctchange_df.sort_values(by=pctchange_df.columns[0], ascending=False)
            for i in range(len(pctchange_df)):
                if pctchange_df.index[i] == "食品饮料":
                    rank = i + 1
                    break
            rank_str = "在申万28个子行业中排名第" + str(rank) + "。"

            top_3_str = "涨幅前三位为" + pctchange_df.index[0] + "（" + format(pctchange_df.iloc[0, 0], ".2f") + "%），" + \
                        pctchange_df.index[1] + "（" + format(pctchange_df.iloc[1, 0], ".2f") + "%），" + \
                        pctchange_df.index[2] + "（" + format(pctchange_df.iloc[2, 0], ".2f") + "%）；"

            sheet_4["A6"].value = pctchange_df.index[0]
            sheet_4["A7"].value = pctchange_df.index[1]
            sheet_4["A8"].value = pctchange_df.index[2]

            return spyl_str + bj_str + szzz_str + active_return_str + rank_str + top_3_str

        def generate_part_2():

            index = sheet_3["B4:B97"].value
            pctchange = sheet_3["D4:D97"].value
            pctchange_df = pd.DataFrame(pctchange, index=index)
            pctchange_df = pctchange_df.sort_values(by=pctchange_df.columns[0], ascending=False)

            top_3_str = "个股方面，涨幅前三位为" + pctchange_df.index[0] + "（" + format(pctchange_df.iloc[0, 0], ".2f") + "%）、" + \
                        pctchange_df.index[1] + "（" + format(pctchange_df.iloc[1, 0], ".2f") + "%）、" + \
                        pctchange_df.index[2] + "（" + format(pctchange_df.iloc[2, 0], ".2f") + "%）；"
            last_3_str = "跌幅前三位为" + pctchange_df.index[-1] + "（" + format(pctchange_df.iloc[-1, 0], ".2f") + "%）、" + \
                         pctchange_df.index[-2] + "（" + format(pctchange_df.iloc[-2, 0], ".2f") + "%）、" + \
                         pctchange_df.index[-3] + "（" + format(pctchange_df.iloc[-3, 0], ".2f") + "%）。"

            sheet_4["B14"].value = pctchange_df.index[0]
            sheet_4["B15"].value = pctchange_df.index[1]
            sheet_4["B16"].value = pctchange_df.index[2]
            sheet_4["B18"].value = pctchange_df.index[-1]
            sheet_4["B19"].value = pctchange_df.index[-2]
            sheet_4["B20"].value = pctchange_df.index[-3]

            return top_3_str + last_3_str

        new_run = generate_part_1() + generate_part_2()
        sheet_4["A2"].value = "一周板块回顾：" + new_run
        workbook.save(file_path)
        print("本周投资提示表格已更改")
        for i in range(len(doc_2.paragraphs)):
            if doc_2.paragraphs[i].text == "1．食品饮料本周观点":
                start_no = i + 1
        run = doc_2.paragraphs[start_no].add_run("一周板块回顾：")
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.bold = True
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run = doc_2.paragraphs[start_no].add_run(new_run)
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        print("一周板块回顾已添加")

    generate_section_1(file_path_3, doc_2)

    # Section 2: 生成本周重要公告及新闻
    print("---------------Step 2: 生成本周重要公告及新闻---------------")
    print("选择需要导入的【每日速递】文件，导入前请确认文件中只包含【本周】的公告及新闻")
    file_path_1 = filedialog.askopenfilename()
    doc_1 = docx.Document(file_path_1)
    print("每日速递文件已导入")

    print("正在执行修改......")

    def generate_section_2(doc_1, doc_2):

        full_text = []
        for p in doc_1.paragraphs:
            full_text.append(p.text)

        useless_contents = ["【申万食品每日速递】",
                            "————————————",
                            "今日行情：",
                            "公司公告&新闻",
                            "申万食品团队",
                            "吕昌|周缘|毕晓静|赵玥"]

        copy = deepcopy(full_text)
        for text in full_text:
            for useless_content in useless_contents:
                if useless_content in text:
                    copy.remove(text)
                else:
                    pass
        null_value = ""
        while null_value in copy:
            copy.remove(null_value)

        pro = ts.pro_api("67386e5aa0f5581d335741b96866313a14fd8bf216902d556dc0fdaf")
        data = pro.stock_basic(exchange='', list_status='L', fields='ts_code,symbol,name,area,industry,list_date')
        data.to_csv("stock_list.csv")
        # data = pd.read_csv("stock_list.csv")
        stocks = []
        for p in copy:
            stock = re.findall(r"【(.+?)】", p)[0]
            stocks.append(stock)
        stock_codes = []
        for stock in stocks:
            for index, row in data.iterrows():
                if row["name"] == stock:
                    stock_code = row["ts_code"]
                    stock_codes.append(stock_code)
        for i in range(len(stocks)):
            strip_str = "【" + stocks[i] + "】"
            copy[i] = copy[i].strip(strip_str)

        stocks = stocks[::-1]
        stock_codes = stock_codes[::-1]
        copy = copy[::-1]

        for i in range(len(doc_2.paragraphs)):
            if doc_2.paragraphs[i].text == "本周重要公告及新闻：":
                start_no = i + 1
            if doc_2.paragraphs[i].text == "表1：肉制品数据":
                end_no = i - 4

        prior_paragraphs = []
        run_1s = []
        run_2s = []
        run_3s = []
        font_1s = []
        font_2s = []
        font_3s = []

        for i in range(len(stocks)):
            prior_paragraphs.append(doc_2.paragraphs[end_no].insert_paragraph_before())
            run_1s.append(prior_paragraphs[i].add_run(str(stocks[i])))
            font_1s.append(run_1s[i].font)
            font_1s[i].name = u"微软雅黑"
            font_1s[i]._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font_1s[i].size = Pt(9)
            font_1s[i].bold = True
            font_1s[i].color.rgb = RGBColor(0x40, 0x40, 0x40)
            run_2s.append(prior_paragraphs[i].add_run("（" + str(stock_codes[i]) + "） "))
            font_2s.append(run_2s[i].font)
            font_2s[i].name = u"微软雅黑"
            font_2s[i]._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font_2s[i].size = Pt(9)
            font_2s[i].bold = True
            font_2s[i].color.rgb = RGBColor(0x40, 0x40, 0x40)
            run_3s.append(prior_paragraphs[i].add_run(str(copy[i])))
            font_3s.append(run_3s[i].font)
            font_3s[i].name = u"微软雅黑"
            font_3s[i]._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font_3s[i].size = Pt(9)
            font_3s[i].color.rgb = RGBColor(0x40, 0x40, 0x40)

        for no in range(start_no, end_no):
            doc_2.paragraphs[no].clear()
        print("本周重要公告及新闻已修改")

    generate_section_2(doc_1, doc_2)

    # Section 3: 生成成本变动回顾
    print("-----------------Step 3: 生成成本变动回顾-----------------")
    print("导入【基础图标更新】表格，导入前请确认各sheet数据已更新，【图3：猪肉价格图】sheet中左侧价格数据已复制至右侧，并生成同比环比数据")
    file_path_5 = filedialog.askopenfilename()
    print("基础图标更新表格已导入")

    print("正在执行修改......")

    def generate_section_3(file_path, doc_2):

        workbook = xw.Book(file_path)
        sheet_1 = workbook.sheets("图1. 新-生鲜乳（周） ")
        sheet_2 = workbook.sheets("图2. 新奶粉统计")
        sheet_3 = workbook.sheets("图3. 猪肉价格图 ")

        date_datetime = datetime.strptime(date, "%Y%m%d") - timedelta(days=9)
        data_1 = sheet_1["A11:D1000"].value
        df_1 = pd.DataFrame(data_1)
        df_1 = df_1.set_index(0)
        sxr = df_1[df_1.index == date_datetime]
        if sxr.iloc[0, 1] >= 0:
            sxr_str_1 = "上升"
        else:
            sxr_str_1 = "下降"
        if sxr.iloc[0, 2] >= 0:
            sxr_str_2 = "上升"
        else:
            sxr_str_2 = "下降"
        part_1 = "主产区生鲜乳平均价为" + format(sxr.iloc[0, 0], ".2f") + "元/公斤，同比" + sxr_str_1 + format(abs(sxr.iloc[0, 1]),
                                                                                               ".2f") + "%，环比" + sxr_str_2 + format(
            abs(sxr.iloc[0, 2]), ".2f") + "%。"

        df_2 = pd.DataFrame(sheet_2["E11:G300"].value, index=sheet_2["A11:A300"].value)
        for i in range(len(df_2)):
            if pd.isnull(df_2.index[i]):
                index_no = i - 1
                break
        if df_2.iloc[index_no, 1] >= 0:
            nf_str_1 = "上升"
        else:
            nf_str_1 = "下降"
        if df_2.iloc[index_no, 2] >= 0:
            nf_str_2 = "上升"
        else:
            nf_str_2 = "下降"
        part_2 = str(df_2.index[index_no].month) + "月份奶粉进口平均价为" + format(df_2.iloc[index_no, 0],
                                                                         ".2f") + "美元/吨，同比" + nf_str_1 + format(
            abs(df_2.iloc[index_no, 1]), ".2%") + "，环比" + nf_str_2 + format(abs(df_2.iloc[index_no, 2]), ".2%") + "。"

        df_3 = pd.DataFrame(sheet_3["I8:Q1000"].value, index=sheet_3["B8:B1000"].value)
        for i in range(len(df_3)):
            if df_3.index[i] == date_datetime:
                index_no = i
                break
        if df_3.iloc[index_no, 3] >= 0:
            zr_str_1 = "上升"
        else:
            zr_str_1 = "下降"
        if df_3.iloc[index_no, 4] >= 0:
            zr_str_2 = "上升"
        else:
            zr_str_2 = "下降"
        if df_3.iloc[index_no, 5] >= 0:
            zr_str_3 = "上升"
        else:
            zr_str_3 = "下降"
        if df_3.iloc[index_no, 6] >= 0:
            zr_str_4 = "上升"
        else:
            zr_str_4 = "下降"
        if df_3.iloc[index_no, 7] >= 0:
            zr_str_5 = "上升"
        else:
            zr_str_5 = "下降"
        if df_3.iloc[index_no, 8] >= 0:
            zr_str_6 = "上升"
        else:
            zr_str_6 = "下降"
        part_3 = "上周仔猪价格" + format(df_3.iloc[index_no, 0], ".2f") + "元/公斤，同比" + zr_str_1 + format(
            abs(df_3.iloc[index_no, 3]), ".2%") + "，环比" + zr_str_2 + format(abs(df_3.iloc[index_no, 4]), ".2%") + "；"
        part_4 = "生猪价格" + format(df_3.iloc[index_no, 1], ".2f") + "元/公斤，同比" + zr_str_3 + format(
            abs(df_3.iloc[index_no, 5]), ".2%") + "，环比" + zr_str_4 + format(abs(df_3.iloc[index_no, 6]), ".2%") + "；"
        part_5 = "猪肉价格" + format(df_3.iloc[index_no, 2], ".2f") + "元/公斤，同比" + zr_str_5 + format(
            abs(df_3.iloc[index_no, 7]), ".2%") + "，环比" + zr_str_6 + format(abs(df_3.iloc[index_no, 8]), ".2%") + "。"

        new_run = part_1 + part_2 + part_3 + part_4 + part_5
        sheet_1["P1"].value = "成本变动回顾：" + new_run
        workbook.save(file_path)
        print("基础图表更新表格已更改")
        for i in range(len(doc_2.paragraphs)):
            if doc_2.paragraphs[i].text == "本周重要公告及新闻：":
                break
        for j in range(1, 21):
            start_no = i - j
            if "成本变动回顾：" in doc_2.paragraphs[start_no].text:
                break
        run = doc_2.paragraphs[start_no].add_run("成本变动回顾：")
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.bold = True
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run = doc_2.paragraphs[start_no].add_run(new_run)
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        print("成本变动回顾已添加")

    generate_section_3(file_path_5, doc_2)

    # Section 4: 生成板块估值水平
    print("-----------------Step 4: 生成板块估值水平-----------------")
    print("请确认【本周投资提示】表格中【动态市盈率PE计算】sheet数据已更新")

    print("正在执行修改......")

    def generate_section_4(file_path, doc_2):

        workbook = xw.Book(file_path)
        sheet = workbook.sheets("动态市盈率PE计算")

        date_datetime = datetime.strptime(date, "%Y%m%d")
        data = sheet["A7:J500"].value
        df = pd.DataFrame(data)
        df = df.set_index(0)
        pe_data = df[df.index == date_datetime]

        if pe_data.iloc[0, 5] >= 0:
            pe_str_1 = "上涨"
        else:
            pe_str_1 = "下跌"
        if pe_data.iloc[0, 8] >= 0:
            pe_str_2 = "上涨"
        else:
            pe_str_2 = "下跌"
        part_1 = "目前食品饮料板块" + str(date_datetime.year) + "年动态PE" + format(pe_data.iloc[0, 0], ".2f") + "x，溢价率" + format(
            pe_data.iloc[0, 4], ".0%") + "，环比" + pe_str_1 + format(abs(pe_data.iloc[0, 5]) * 100, ".1f") + "pct；"
        part_2 = "白酒动态PE" + format(pe_data.iloc[0, 1], ".2f") + "x，溢价率" + format(pe_data.iloc[0, 7],
                                                                                 ".0%") + "，环比" + pe_str_2 + format(
            abs(pe_data.iloc[0, 8]) * 100, ".1f") + "pct。"

        new_run = part_1 + part_2
        sheet["A2"].value = new_run
        workbook.save(file_path)
        print("本周投资提示表格-动态市盈率sheet已更改")
        for i in range(len(doc_2.paragraphs)):
            if doc_2.paragraphs[i].text == "本周重要公告及新闻：":
                break
        for j in range(1, 21):
            start_no = i - j
            if "板块估值水平：" in doc_2.paragraphs[start_no].text:
                break
        run = doc_2.paragraphs[start_no].add_run("板块估值水平：")
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.bold = True
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run = doc_2.paragraphs[start_no].add_run(new_run)
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        print("板块估值水平已添加")

    generate_section_4(file_path_3, doc_2)

    # Section 5: 生成市场表现
    print("-----------------Step 5: 生成市场表现-----------------")
    print("导入【市场表现】表格，导入前请确认数据提取日期已更改为本周最后一个交易日且数据已更新")
    file_path_6 = filedialog.askopenfilename()
    print("市场表现表格已导入")

    def generate_section_5(file_path, doc_2):

        workbook = xw.Book(file_path)
        sheet_1 = workbook.sheets("子行业超额收益")

        start_date = datetime.strptime(date, "%Y%m%d") - timedelta(days=4)
        end_date = datetime.strptime(date, "%Y%m%d")

        print("正在检查数据提取日期是否已更新......")
        if sheet_1["A1"].value == end_date:
            pass
        else:
            print("【子行业超额收益】sheet日期未更新，请更新后重试！")
            exit()

        print("正在执行修改......")

        part_1 = str(start_date.year)[2:4] + "/" + str(start_date.month) + "/" + str(start_date.day) + "-" + str(
            end_date.year)[2:4] + "/" + str(end_date.month) + "/" + str(end_date.day) + "，"
        data = sheet_1["A8:G45"].value
        df = pd.DataFrame(data)
        df = df.set_index(0)
        if df[df.index == "食品饮料"].iloc[0, 1] >= 0:
            str_1 = "跑赢"
        else:
            str_1 = "跑输"
        part_2 = "食品饮料行业" + str_1 + "申万A指" + format(abs(df[df.index == "食品饮料"].iloc[0, 1]) * 100,
                                                    ".1f") + "个百分点，子行业表现依次为："

        sub_industries = ["调味发酵品", "白酒", "啤酒", "葡萄酒", "肉制品", "乳品"]
        sub_data = []
        for sub_industry in sub_industries:
            sub_data.append(df[df.index == sub_industry].iloc[0, 1])
        sub_dict = dict(zip(sub_industries, sub_data))
        ranks = sorted(sub_dict, key=sub_dict.get, reverse=True)
        strs = []
        active_returns = []
        for rank in ranks:
            active_returns.append(format(abs(sub_dict[rank]) * 100, ".1f"))
            if sub_dict[rank] >= 0:
                strs.append("跑赢")
            else:
                strs.append("跑输")
        part_3 = "{0[0]}（{1[0]}申万A指{2[0]}个百分点）、{0[1]}（{1[1]}申万A指{2[1]}个百分点）、{0[2]}（{1[2]}申万A指{2[2]}个百分点）、{0[3]}（{1[3]}申万A指{2[3]}个百分点）、{0[4]}（{1[4]}申万A指{2[4]}个百分点）、{0[5]}（{1[5]}申万A指{2[5]}个百分点）。".format(
            ranks, strs, active_returns)

        new_run = part_1 + part_2 + part_3
        sheet_1["L5"].value = ranks[0]
        sheet_1["L6"].value = ranks[1]
        sheet_1["L7"].value = ranks[2]
        sheet_1["L8"].value = ranks[3]
        sheet_1["L9"].value = ranks[4]
        sheet_1["L10"].value = ranks[5]
        sheet_1["I14"].value = new_run
        workbook.save(file_path)
        print("市场表现表格已更改")

        for i in range(len(doc_2.paragraphs)):
            if doc_2.paragraphs[i].text == "2．食品饮料各板块市场表现":
                start_no = i + 1
        run = doc_2.paragraphs[start_no].add_run(new_run)
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        print("市场表现已添加")

        sheet_2 = workbook.sheets("公司新0517")

        print("正在修改表9、表10......")

        for character_1 in characters:
            col_no_1 = character_1 + "5"
            if sheet_2[col_no_1].value == end_date:
                break

        active_returns_1 = sheet_2[character_1 + "7:" + character_1 + "40"].value

        doc_2.tables[17].cell(1, -1).text = sheet_2[character_1 + "3"].value
        doc_2.tables[17].cell(1, -1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = doc_2.tables[17].cell(1, -1).paragraphs[0].runs[0]
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(9)
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)

        for i in range(2, 35):
            # print(doc_2.tables[17].cell(i, -1).text)
            doc_2.tables[17].cell(i, -1).text = format(active_returns_1[i - 2], ".1%")
            doc_2.tables[17].cell(i, -1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[17].cell(i, -1).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        active_returns_2 = sheet_1["O31:O39"].value

        doc_2.tables[15].cell(1, -1).text = sheet_2[character_1 + "3"].value
        doc_2.tables[15].cell(1, -1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = doc_2.tables[15].cell(1, -1).paragraphs[0].runs[0]
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(9)
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)

        for i in range(2, 11):
            # print(doc_2.tables[17].cell(i, -1).text)
            doc_2.tables[15].cell(i, -1).text = format(active_returns_2[i - 2], ".1%")
            doc_2.tables[15].cell(i, -1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[15].cell(i, -1).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        print("表9、表10已更改")

    generate_section_5(file_path_6, doc_2)

    # Section 6: 生成食品大类价格表格
    print("-----------------Step 6: 生成食品大类价格表格-----------------")
    print("导入【食品大类价格】表格，导入前请确认数据已更新")
    file_path_7 = filedialog.askopenfilename()
    print("食品大类价格表格已导入")

    print("正在执行修改......")

    def generate_section_6(file_path, doc_2):

        workbook = xw.Book(file_path)
        sheet_1 = workbook.sheets("输出页")

        def generate_cols(sheet, sheet_range, table_no):

            data = sheet[sheet_range].value
            df = pd.DataFrame(data)
            for i in range(1, len(df) + 1):
                # modify col 1
                # print(doc_2.tables[6].cell(i, 1).text)
                # print("{0}/{1}/{2}".format(df.iloc[i-1, 1].year, df.iloc[i-1, 1].month, df.iloc[i-1, 1].day))
                doc_2.tables[table_no].cell(i, 1).text = "{0}/{1}/{2}".format(df.iloc[i - 1, 1].year,
                                                                              df.iloc[i - 1, 1].month,
                                                                              df.iloc[i - 1, 1].day)
                doc_2.tables[table_no].cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 1).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 2
                exceptions = ["生猪存栏", "能繁母猪", "GDT成交量", "GDT成交价", "进口奶粉金额", "进口奶粉数量", "进口液奶金额", "进口液奶数量", "进口大麦数量"]
                if df.iloc[i - 1, 0] in exceptions or sheet_range == "A71:H86" or sheet_range == "A91:H100":
                    doc_2.tables[table_no].cell(i, 2).text = format(df.iloc[i - 1, 2], ".0f")
                else:
                    doc_2.tables[table_no].cell(i, 2).text = format(df.iloc[i - 1, 2], ".2f")
                doc_2.tables[table_no].cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 2).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 3
                if df.iloc[i - 1, 3] == "-":
                    doc_2.tables[table_no].cell(i, 3).text = df.iloc[i - 1, 3]
                else:
                    doc_2.tables[table_no].cell(i, 3).text = format(df.iloc[i - 1, 3], ".2%")
                doc_2.tables[table_no].cell(i, 3).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 3).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 4
                if df.iloc[i - 1, 4] == "-":
                    doc_2.tables[table_no].cell(i, 4).text = df.iloc[i - 1, 4]
                else:
                    doc_2.tables[table_no].cell(i, 4).text = format(df.iloc[i - 1, 4], ".2%")
                doc_2.tables[table_no].cell(i, 4).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 4).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 5
                if df.iloc[i - 1, 5] == "-":
                    doc_2.tables[table_no].cell(i, 5).text = df.iloc[i - 1, 5]
                else:
                    doc_2.tables[table_no].cell(i, 5).text = format(df.iloc[i - 1, 5], ".2%")
                doc_2.tables[table_no].cell(i, 5).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 5).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 6
                if pd.isnull(df.iloc[i - 1, 6]):
                    doc_2.tables[table_no].cell(i, 6).text = ""
                else:
                    doc_2.tables[table_no].cell(i, 6).text = df.iloc[i - 1, 6]
                doc_2.tables[table_no].cell(i, 6).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 6).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                # modify col 7
                doc_2.tables[table_no].cell(i, 7).text = df.iloc[i - 1, 7]
                doc_2.tables[table_no].cell(i, 7).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[table_no].cell(i, 7).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        sheet_range_1 = "A4:H12"
        table_no_1 = 6
        sheet_range_2 = "A18:H32"
        table_no_2 = 7
        sheet_range_3 = "A37:H46"
        table_no_3 = 8
        sheet_range_4 = "A51:H54"
        table_no_4 = 9
        sheet_range_5 = "A60:H66"
        table_no_5 = 10
        sheet_range_6 = "A71:H86"
        table_no_6 = 11
        sheet_range_7 = "A91:H100"
        table_no_7 = 12
        sheet_range_8 = "A106:H107"
        table_no_8 = 13

        generate_cols(sheet_1, sheet_range_1, table_no_1)
        generate_cols(sheet_1, sheet_range_2, table_no_2)
        generate_cols(sheet_1, sheet_range_3, table_no_3)
        generate_cols(sheet_1, sheet_range_4, table_no_4)
        generate_cols(sheet_1, sheet_range_5, table_no_5)
        generate_cols(sheet_1, sheet_range_6, table_no_6)
        generate_cols(sheet_1, sheet_range_7, table_no_7)
        generate_cols(sheet_1, sheet_range_8, table_no_8)
        print("表1—8已更改")

    generate_section_6(file_path_7, doc_2)

    # Section 7: 生成盈利预测表格
    print("-----------------Step 7: 生成盈利预测表格-----------------")
    print("导入【盈利预测】表格，导入前请确认数据已更新")
    file_path_8 = filedialog.askopenfilename()
    print("盈利预测表格已导入")

    print("正在执行修改......")

    def generate_section_7(file_path, doc_2):

        workbook = xw.Book(file_path)
        sheet_1 = workbook.sheets("盈利预测 (报告版本)")

        data = sheet_1["A4:P30"].value
        df = pd.DataFrame(data)

        date_datetime = datetime.strptime(date, "%Y%m%d")
        doc_2.tables[-1].cell(0, 1).text = "股价（元）\n（{}-{}）".format(format(date_datetime.month, "0>2d"), format(date_datetime.day, "0>2d"))
        run = doc_2.tables[-1].cell(0, 1).paragraphs[0].runs[0]
        # print(run.text)
        font = run.font
        font.name = u"微软雅黑"
        font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        font.size = Pt(9)
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)

        for i in range(2, len(df) + 2):
            # modify col 1
            doc_2.tables[-1].cell(i, 0).text = df.iloc[i - 2, 0]
            doc_2.tables[-1].cell(i, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[-1].cell(i, 0).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 2
            doc_2.tables[-1].cell(i, 1).text = format(df.iloc[i - 2, 1], ".2f")
            doc_2.tables[-1].cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[-1].cell(i, 1).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 3 - 6
            for j in range(2, 6):
                doc_2.tables[-1].cell(i, j).text = format(df.iloc[i - 2, j + 1], ".1f")
                doc_2.tables[-1].cell(i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[-1].cell(i, j).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 7 - 10
            for k in range(6, 10):
                doc_2.tables[-1].cell(i, k).text = format(df.iloc[i - 2, k + 1], ".0%")
                doc_2.tables[-1].cell(i, k).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[-1].cell(i, k).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 11 - 13
            for l in range(10, 13):
                doc_2.tables[-1].cell(i, l).text = format(df.iloc[i - 2, l + 1], ".0f")
                doc_2.tables[-1].cell(i, l).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = doc_2.tables[-1].cell(i, l).paragraphs[0].runs[0]
                # print(run.text)
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 14
            doc_2.tables[-1].cell(i, 13).text = df.iloc[i - 2, 14]
            doc_2.tables[-1].cell(i, 13).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[-1].cell(i, 13).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            # modify col 15
            doc_2.tables[-1].cell(i, 14).text = format(df.iloc[i - 2, 15], ".0f")
            doc_2.tables[-1].cell(i, 14).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = doc_2.tables[-1].cell(i, 14).paragraphs[0].runs[0]
            # print(run.text)
            font = run.font
            font.name = u"微软雅黑"
            font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
            font.size = Pt(9)
            font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        print("表15已更改")

    generate_section_7(file_path_8, doc_2)

    # Section 8: 修改日期
    print("-----------------Step 8: 修改日期-----------------")
    print("正在执行修改......")

    def generate_section_8(doc_2):

        date_datetime = datetime.strptime(date, "%Y%m%d")

        # for i in range(len(doc_2.paragraphs)):
        #     if doc_2.paragraphs[i].text == "表9：去年及今年食品饮料各子行业相对于申万A指的超额收益情况":
        #         start_no = i + 1
        # doc_2.paragraphs[start_no].text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
        #                                                                            format(date_datetime.month, "0>2d"),
        #                                                                            format(date_datetime.day, "0>2d"))
        # run = doc_2.paragraphs[start_no].runs[0]
        # # print(run.text)
        # font = run.font
        # font.name = u"微软雅黑"
        # font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        # font.size = Pt(9)
        # font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # for i in range(len(doc_2.paragraphs)):
        #     if doc_2.paragraphs[i].text == "表10：食品饮料跟踪公司相对于申万A指的超额收益情况":
        #         start_no = i + 1
        # doc_2.paragraphs[start_no].text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
        #                                                                            format(date_datetime.month, "0>2d"),
        #                                                                            format(date_datetime.day, "0>2d"))
        # run = doc_2.paragraphs[start_no].runs[0]
        # # print(run.text)
        # font = run.font
        # font.name = u"微软雅黑"
        # font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        # font.size = Pt(9)
        # font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # for i in range(len(doc_2.paragraphs)):
        #     if doc_2.paragraphs[i].text == "表11：本周重要股东增减持明细":
        #         start_no = i + 1
        # doc_2.paragraphs[start_no].text = "资料来源：Wind，申万宏源研究（数据截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
        #                                                                              format(date_datetime.month,
        #                                                                                     "0>2d"),
        #                                                                              format(date_datetime.day, "0>2d"))
        # run = doc_2.paragraphs[start_no].runs[0]
        # # print(run.text)
        # font = run.font
        # font.name = u"微软雅黑"
        # font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        # font.size = Pt(9)
        # font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # for i in range(len(doc_2.paragraphs)):
        #     if doc_2.paragraphs[i].text == "表15：食品饮料重点公司盈利预测表":
        #         start_no = i + 1
        # doc_2.paragraphs[start_no].text = "资料来源：Wind，申万宏源研究（数据截至{0}/{1}/{2}）；注：千禾味业21年预测净利润为万得一致预期".format(
        #     str(date_datetime.year)[2:], format(date_datetime.month, "0>2d"), format(date_datetime.day, "0>2d"))
        # run = doc_2.paragraphs[start_no].runs[0]
        # # print(run.text)
        # font = run.font
        # font.name = u"微软雅黑"
        # font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        # font.size = Pt(9)
        # font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # references = [
        #     doc_2.tables[3].cell(2, 0),  # 图1
        #     doc_2.tables[3].cell(2, 2),  # 图2
        #     doc_2.tables[4].cell(2, 0),  # 图3
        #     doc_2.tables[4].cell(2, 2),  # 图4
        #     doc_2.tables[5].cell(2, 0),  # 图5
        #     doc_2.tables[5].cell(2, 2),  # 图6
        #     doc_2.tables[14].cell(2, 0),  # 图7
        #     doc_2.tables[16].cell(2, 0),  # 图8
        #     doc_2.tables[16].cell(2, 2),  # 图9
        #     doc_2.tables[19].cell(7, 0),  # 表12
        #     doc_2.tables[20].cell(3, 0),  # 表13
        #     doc_2.tables[22].cell(2, 0),  # 图10
        #     doc_2.tables[22].cell(2, 2),  # 图11
        #     doc_2.tables[23].cell(2, 0),  # 图12
        #     doc_2.tables[23].cell(2, 2),  # 图13
        #     doc_2.tables[24].cell(2, 0),  # 图14
        #     doc_2.tables[24].cell(2, 3),  # 图15
        #     doc_2.tables[25].cell(2, 0),  # 图16
        #     doc_2.tables[25].cell(2, 2),  # 图17
        #     doc_2.tables[26].cell(2, 0)  # 图18
        # ]
        #
        # for i in range(0, 2):
        #     references[i].text = "资料来源：奶业协会，申万宏源研究（截至{0}年{1}月）".format(str(date_datetime.year)[2:], date_datetime.month)
        #     run = references[i].paragraphs[0].runs[0]
        #     # print(run.text)
        #     font = run.font
        #     font.name = u"微软雅黑"
        #     font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        #     font.size = Pt(9)
        #     font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # references[2].text = "资料来源：农业部，申万宏源研究（截至{0}年{1}月）".format(str(date_datetime.year)[2:], date_datetime.month)
        # run = references[2].paragraphs[0].runs[0]
        # # print(run.text)
        # font = run.font
        # font.name = u"微软雅黑"
        # font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        # font.size = Pt(9)
        # font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # for i in [6, 9, 10]:
        #     references[i].text = "资料来源：Wind，申万宏源研究（数据截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
        #                                                                     format(date_datetime.month, "0>2d"),
        #                                                                     format(date_datetime.day, "0>2d"))
        #     run = references[i].paragraphs[0].runs[0]
        #     # print(run.text)
        #     font = run.font
        #     font.name = u"微软雅黑"
        #     font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        #     font.size = Pt(9)
        #     font.color.rgb = RGBColor(0x00, 0x32, 0x96)
        #
        # for i in [7, 8, 11, 12, 13, 14, 15, 16, 17, 18, 19]:
        #     references[i].text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
        #                                                                   format(date_datetime.month, "0>2d"),
        #                                                                   format(date_datetime.day, "0>2d"))
        #     run = references[i].paragraphs[0].runs[0]
        #     # print(run.text)
        #     font = run.font
        #     font.name = u"微软雅黑"
        #     font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
        #     font.size = Pt(9)
        #     font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        for i in range(len(doc_2.paragraphs)):
            if "资料来源：Wind，申万宏源研究（截至" in doc_2.paragraphs[i].text:
                doc_2.paragraphs[i].text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
                                                                                    format(date_datetime.month, "0>2d"),
                                                                                    format(date_datetime.day, "0>2d"))
                run = doc_2.paragraphs[i].runs[0]
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            elif "资料来源：Wind，申万宏源研究（数据截至" in doc_2.paragraphs[i].text and "注：千禾味业21年预测净利润为万得一致预期" not in \
                    doc_2.paragraphs[i].text:
                doc_2.paragraphs[i].text = "资料来源：Wind，申万宏源研究（数据截至{0}/{1}/{2}）".format(str(date_datetime.year)[2:],
                                                                                      format(date_datetime.month,
                                                                                             "0>2d"),
                                                                                      format(date_datetime.day, "0>2d"))
                run = doc_2.paragraphs[i].runs[0]
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            elif "资料来源：Wind，申万宏源研究（数据截至" in doc_2.paragraphs[i].text and "注：千禾味业21年预测净利润为万得一致预期" in doc_2.paragraphs[
                i].text:
                doc_2.paragraphs[i].text = "资料来源：Wind，申万宏源研究（数据截至{0}/{1}/{2}）；注：千禾味业21年预测净利润为万得一致预期".format(
                    str(date_datetime.year)[2:], format(date_datetime.month, "0>2d"), format(date_datetime.day, "0>2d"))
                run = doc_2.paragraphs[i].runs[0]
                font = run.font
                font.name = u"微软雅黑"
                font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                font.size = Pt(9)
                font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        for i in range(len(doc_2.tables)):
            if len(doc_2.tables[i].rows) >= 3 and len(doc_2.tables[i].columns) >= 3:
                if "资料来源：Wind，申万宏源研究（截至" in doc_2.tables[i].cell(2, 0).text:
                    doc_2.tables[i].cell(2, 0).text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(
                        str(date_datetime.year)[2:],
                        format(date_datetime.month, "0>2d"),
                        format(date_datetime.day, "0>2d"))
                    run = doc_2.tables[i].cell(2, 0).paragraphs[0].runs[0]
                    font = run.font
                    font.name = u"微软雅黑"
                    font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                    font.size = Pt(9)
                    font.color.rgb = RGBColor(0x00, 0x32, 0x96)
                if "资料来源：Wind，申万宏源研究（截至" in doc_2.tables[i].cell(2, 2).text:
                    doc_2.tables[i].cell(2, 2).text = "资料来源：Wind，申万宏源研究（截至{0}/{1}/{2}）".format(
                        str(date_datetime.year)[2:],
                        format(date_datetime.month, "0>2d"),
                        format(date_datetime.day, "0>2d"))
                    run = doc_2.tables[i].cell(2, 2).paragraphs[0].runs[0]
                    font = run.font
                    font.name = u"微软雅黑"
                    font._element.rPr.rFonts.set(qn("w:eastAsia"), u"微软雅黑")
                    font.size = Pt(9)
                    font.color.rgb = RGBColor(0x00, 0x32, 0x96)

        print("已完成修改")

    generate_section_8(doc_2)

    print("-----------------周报修改已完成，请选择保存路径-----------------")
    file_path_4 = filedialog.asksaveasfilename(filetypes=[("DOCX", ".docx")])
    doc_2.save(str(file_path_4) + ".docx")
    print("保存成功！")


generate_sections()
